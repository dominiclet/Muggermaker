from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy


# opens the default template document
document = Document('default.docx')

# copy default table
table = document.tables[0]

def addcase():
    tbl = table._tbl
    newtable = deepcopy(tbl)
    paragraph = document.add_paragraph()
    paragraph._p.addnext(newtable)

while True:
    try:
        # request for word file to parse
        read_file = input("Reading List: ")
        read_doc = Document(read_file)
        break
    except:
        print("Document does not exist! (Try moving document into current file location)")

# choose output file's name
outfile = input("Output file's name: ")

# iterate over paragraphs in doc
paragraphs = read_doc.paragraphs
cases = []
for para in paragraphs:
    if para.runs:
        if para.runs[0].italic == True:
            case = para.text
            cases.append(case)

# iterate over list of cases and put it in muggers
count = 0
for case in cases:
    cell = document.tables[count].cell(0,0)
    # text
    cell.text = case
    # bold text
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    # center text
    format = cell.paragraphs[0].paragraph_format
    format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # add another empty table
    addcase()
    count += 1

document.save(outfile+'.docx')
